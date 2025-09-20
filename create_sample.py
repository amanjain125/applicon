import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'app'))

# Create a simple DOCX file
from docx import Document

def create_sample_docx():
    print("=== Creating Sample DOCX File ===")
    
    # Create document
    doc = Document()
    
    # Add content
    doc.add_heading('John Doe', 0)
    doc.add_heading('Data Scientist', level=1)
    
    doc.add_paragraph('Email: john.doe@example.com')
    doc.add_paragraph('Phone: (555) 123-4567')
    doc.add_paragraph('LinkedIn: linkedin.com/in/johndoe')
    
    doc.add_heading('SUMMARY', level=2)
    doc.add_paragraph('Experienced data scientist with 3 years of experience in Python and machine learning.')
    
    doc.add_heading('EXPERIENCE', level=2)
    doc.add_paragraph('Data Scientist | Tech Solutions Inc. | jan2020 - Present')
    doc.add_paragraph('- Developed machine learning models using Python and TensorFlow')
    doc.add_paragraph('- Created dashboards with Power BI')
    
    # Save document
    doc.save("sample_resume.docx")
    print("Sample DOCX file created: sample_resume.docx")
    
    # Create sample JD
    with open("sample_jd.txt", "w") as f:
        f.write("""Data Scientist Position

We are looking for a Data Scientist with experience in Python and machine learning.

Requirements:
- 2+ years of experience with Python
- Experience with machine learning frameworks
- Knowledge of data visualization tools

Nice to have:
- Experience with Power BI
- Cloud platform experience
""")
    
    print("Sample JD file created: sample_jd.txt")

if __name__ == "__main__":
    # Check if python-docx is installed
    try:
        from docx import Document
        create_sample_docx()
    except ImportError:
        print("python-docx is not installed. Please install it with: pip install python-docx")